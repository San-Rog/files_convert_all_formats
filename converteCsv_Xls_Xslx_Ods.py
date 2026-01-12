import streamlit as st
import pandas as pd
import os
import io
import csv
import time
import yaml
import toml
import zipfile
import regex as re
import unidecode
from io import BytesIO
from io import StringIO
import openpyxl
import xlrd
import odf
import textwrap
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from fpdf import FPDF 
import streamlit.components.v1 as components
import warnings
warnings.filterwarnings("ignore", category=UserWarning)
from pyexcel_xlsx import get_data
from pyexcel_ods3 import save_data
from collections import OrderedDict
import pyexcel
from collections import Counter
import locale

class messages():
    def __init__(self, *args):
        self.fileTmp = args[0]
        self.suffix = args[1]
        self.nFiles = args[2]
        if self.nFiles == 1:
            self.fileFinal = f'zipado_isolado_{self.suffix}.zip'
        else:
            self.fileFinal = f'zipado_múltiplos_{self.suffix}.zip'
        if None not in args:
            self.mensResult()
    
    def mensResult(self):
        exclRep = st.session_state[replDown[0]]        
        if exclRep: 
            arrayFile = ['arquivo não redundante', 'arquivos não redundantes']
        else:
            arrayFile = ['arquivo repetido', 'arquivos com e sem redundância']
        if self.nFiles <= 1:
            exprFile = [arrayFile[0], 'baixá-lo', 'abri-lo']
        else:
            exprFile = [arrayFile[1], 'baixá-los', 'abri-los']
        if self.suffix in ['tsv', 'yaml', 'json', 'toml', 'txt']:
            mensStr = f':blue[**{self.fileFinal}**] com  ***{self.nFiles} {exprFile[0]}***. Para {exprFile[1]}, ' \
            f'clique no botão ao lado 👉. (Utilize **Bloco de Notas** ou aplicativo similar para {exprFile[2]}.)'
        else:
            mensStr = f':blue[**{self.fileFinal}**] com  ***{self.nFiles} {exprFile[0]}***. Para {exprFile[1]}, ' \
            'clique no botão ao lado 👉.' 
        mensStr = textwrap.fill(mensStr, width=80)
        colMens, colZip = st.columns([21, 3], width='stretch', vertical_alignment='center')
        colMens.success(mensStr, icon='✔️',  width='stretch')                              
        with open(self.fileTmp, "rb") as file:
            buttDown = colZip.download_button(label='',
                                              data=file,
                                              file_name=self.fileFinal,
                                              mime='application/zip', 
                                              icon=':material/download:', 
                                              width='stretch', 
                                              key='buttDown', 
                                              help='Grava o arquivo zipado na pasta Download.')
     
    @st.dialog('⚠️ Falha no app❗')
    def mensOperation(self, str):
        st.markdown(f'{str} Entre em contato com o administrador da ferramenta!')
    
class acessories():
    def __init__(self, *args):
        try:
            locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
        except locale.Error:
            locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')
        
    def valueMoney(self, value):
        valFormat = locale.currency(value, symbol=False, grouping=True)[:-3]
        return valFormat
        
class downOrDfFiles():
    def __init__(self, *args):
        self.files = args[0]
        self.index = args[1]
        self.engine = args[2]
        self.ext = args[3]
        self.opt = args[4]
        self.typeFile = args[5]
        self.typeExt = args[6]
        self.filesZip = []
        self.nFiles = 0
        try:
            self.pos = self.typeExt.index(self.typeFile)
        except:
            self.pos = None
        if self.opt in [-1, -2, -3]:
            self.filesAllDf()
        elif self.opt == -4:
            pass
        else:
            buttSel = True
            match self.pos:
                case 0: 
                    if self.index == 0:
                        if self.opt in [0, 1]:
                            self.csvXlsx() 
                        elif self.opt == 2:
                            self.csvHtml()
                    elif self.index == 1:
                        self.csvOds()
                    elif self.index == 2:
                        self.csvTsv()
                    elif self.index == 3:
                        self.csvDocx()     
                    elif self.index == 4:
                        self.csvYaml()
                    elif self.index == 5:
                        self.csvJson()
                    elif self.index == 6:
                        self.csvXhtml()
                    elif self.index == 7:
                        self.csvToml()
                    elif self.index == 8:
                        self.csvTxt()
                    elif self.index == 9:
                        self.csvPdf()
                    else: 
                         buttSel = False
                case 1:
                    if self.index == 0:
                        if self.opt == 0:
                            self.xlsXlsxOdsManyFormats()
                        elif self.opt == 1:
                            self.xlsXlsxAround()
                        elif self.opt == 2:
                            self.engine = allEngines[1]
                            self.xlsXlsxOdsHtml()                            
                    elif self.index == 1:
                        self.xlsXslxOds()
                    elif self.index == 2:
                        self.engine = allEngines[1]
                        self.xlsXlsxOdsManyFormats()  
                    elif self.index == 3:
                        self.xlsXlsxOdsDocx()
                    elif self.index == 4:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 5:
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 6:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 7:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 8:
                        self.engine = allEngines[1]
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 9:
                        self.xlsXlsxOdsPdf()
                    else:
                        buttSel = False
                case 2: 
                    if self.index == 0:
                        if self.opt == 0:
                            self.xlsXlsxOdsManyFormats() 
                        elif self.opt == 1:
                           self.xlsXlsxAround() 
                        elif self.opt == 2:
                            self.engine = allEngines[0]
                            self.xlsXlsxOdsHtml()
                    elif self.index == 1:
                        self.xlsXslxOds()
                    elif self.index == 2:
                        self.engine = allEngines[0]
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 3:
                        self.xlsXlsxOdsDocx()
                    elif self.index == 4:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 5:
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 6:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 7:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 8:
                        self.engine = allEngines[0]
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 9:
                        self.xlsXlsxOdsPdf()                    
                case 3:
                    if self.index == 0:
                        if self.opt in [0, 1]:
                            self.odsXlsXlsx()
                        elif self.opt == 2:
                            self.engine = allEngines[2]
                            self.xlsXlsxOdsHtml()
                    if self.index == 1:
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 2:
                        self.engine = allEngines[2]
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 3:
                        self.xlsXlsxOdsDocx()
                    elif self.index == 4:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 5:
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 6:
                        self.xlsXlsxOdsYaml()                    
                    elif self.index == 7:
                        self.xlsXlsxOdsYaml()
                    elif self.index == 8:
                        self.engine = allEngines[2]
                        self.xlsXlsxOdsManyFormats()
                    elif self.index == 9:
                        self.xlsXlsxOdsPdf()                    
            if all([self.opt is not None, buttSel]):
                self.nameZip = f'arquivo_all_{self.ext}.zip'
                self.downZip()
                if os.path.getsize(self.nameZip) > 0:
                    messages(self.nameZip, self.ext, self.nFiles) 
        
    def xlsXlsxAround(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                dfAllDict = allDfs[f]   
                self.dfAll = dfAllDict[0]
            except: 
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
            for name, df in self.dfAll.items():
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                self.df = df.fillna('')
                self.df = self.df.astype(str)
                try:
                    df.to_excel(self.fileOut, index=False, na_rep='')
                except: 
                    df.to_excel(self.fileOut, index=False, engine=self.engine, na_rep='')
                self.bytesFiles()
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_aba_por_aba.{self.ext}'
            with pd.ExcelWriter(self.fileOut, engine='openpyxl') as writer:
                for name, df in self.dfAll.items():
                    df.to_excel(writer, sheet_name=name, index=False) 
            self.bytesFiles()
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_abas_sobrepostas.{self.ext}'
            self.filesUniqueFile(1)    
    
    def xlsXlsxOdsPdf(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                dfAllDict = allDfs[f]   
                self.dfAll = dfAllDict[0]
            except: 
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
            for name, df in self.dfAll.items():
                self.fileOut = f'{self.nameFile}_{f+1}_{name}.{self.ext}'
                self.df = df.fillna('')
                self.df = self.df.astype(str)
                self.prepairePdf()
                self.bytesFiles()
            self.fileOut = f'{self.nameFile}_{f+1}.{self.ext}'
            self.filesUniqueFile(6)
    
    def csvOds(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1] 
            self.sep = file[-1]
            self.fileOut = f'{self.nameFile}_new.csv'
            self.csvCsv(0)
            self.fileCsv = self.fileOut
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            pyexcel.save_as(file_name=self.fileCsv, dest_file_name=self.fileOut)
            self.bytesFiles()            
    
    def csvHtml(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            self.df = pd.read_csv(self.fileOut).fillna('')
            self.renameHead()
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            htmlTable = self.df.to_html()
            with open(self.fileOut, 'w', encoding='utf-8-sig') as f:
                f.write(htmlTable)
            self.bytesFiles()
            
    def xlsXlsxOdsHtml(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None, engine=self.engine)
            except:
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
            for name, df in self.dfAll.items():
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                self.df = df.fillna('')
                self.df = self.df.astype(str)
                self.renameHead()
                try:
                    htmlStr = df.to_html(index=False, border=1, classes='dataframe')
                except:
                    htmlStr = df.to_html()
                with open(self.fileOut, 'w', encoding='utf-8') as f:
                    f.write(htmlStr)
                self.bytesFiles()
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            self.filesUniqueFile(2)
            
    def xlsXlsxOdsYaml(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                dfAllDict = allDfs[f]   
                self.dfAll = dfAllDict[0]
            except: 
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
            for name, df in self.dfAll.items():
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                self.df = df.fillna('')
                self.df = self.df.astype(str)
                try:
                    self.renameHead()
                except: 
                    pass
                if self.index == 0:
                    htmlStr = self.df.to_html()
                    with open(self.fileOut, 'w', encoding='utf-8') as f:
                        f.write(htmlStr)
                elif self.index == 4:
                    yamlData = {}
                    yamlData[name] = df.to_dict(orient='records')
                    with open(self.fileOut, 'w', encoding='utf-8') as outfile:
                        yaml.dump(yamlData, outfile, sort_keys=False, indent=4, allow_unicode=True)
                elif self.index == 6:
                    htmlTable = df.to_html(index=True, border=1, classes='dataframe', na_rep='', 
                                           justify='center', show_dimensions=True)
                    external = configExternal(None)
                    xhtmlContent = external.includeXhtml(htmlTable)
                    self.fileOut = f'{self.nameFile}.{self.ext}'
                    with open(self.fileOut, 'w', encoding='utf-8') as f:
                        f.write(xhtmlContent)                
                elif self.index == 7:
                    tomData = {}
                    tomData[name] = df.to_dict(orient='records')
                    with open(self.fileOut, 'w', encoding='utf-8') as outfile:
                        yaml.dump(tomData, outfile, sort_keys=False, indent=4, allow_unicode=True)
                self.bytesFiles()
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            self.filesUniqueFile(3)
            
    def xlsXlsxOdsDocx(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                dfAllDict = allDfs[f]   
                self.dfAll = dfAllDict[0]
            except: 
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
            self.docs = []
            for name, df in self.dfAll.items():
                doc = Document()
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                self.df = df.fillna('')
                #self.renameHead()
                doc.add_heading(f'Tabela da Aba: {name}', level=1)
                table = doc.add_table(rows=1, cols=len(self.df.columns))
                table.style = 'Table Grid' 
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(self.df.columns):
                    hdr_cells[i].text = str(col)
                for index, row in self.df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                doc.add_page_break()
                doc.save(self.fileOut)
                self.docs.append(self.fileOut)
                self.bytesFiles()
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            self.filesUniqueFile(4)
    
    def odsXlsXlsx(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                dfAllDict = allDfs[f]   
                self.dfAll = dfAllDict[0]
            except: 
                self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
            for name, df in self.dfAll.items():
                self.fileOut = f'{self.nameFile}_prov_{name}.ods'
                with pd.ExcelWriter(self.fileOut, engine="odf") as writer:
                    df.to_excel(writer, index=False, sheet_name=name)
                self.fileOds = self.fileOut
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                pyexcel.save_as(file_name=self.fileOds, dest_file_name=self.fileOut)
                self.bytesFiles()   
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            self.filesUniqueFile(7)
            
    def xlsXslxOds(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            try:
                dfAllDict = allDfs[f]   
                self.df = dfAllDict[0]
            except: 
                self.df = pd.read_excel(self.dataFile, sheet_name=None)
            self.sheets = list(self.df.keys())
            self.fileOut = f'{self.nameFile}_prov.xls'
            self.createAllPlan()
            self.df = pd.read_excel(self.fileOut, sheet_name=None)
            for name, df in self.df.items():
                self.fileOut = f'{self.nameFile}_{name}.xlsx'
                with pd.ExcelWriter(self.fileOut, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name=name, index=False)            
                dataXlsx = get_data(self.fileOut) 
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                save_data(self.fileOut, dataXlsx)                
                self.bytesFiles() 
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_aba_por_aba.{self.ext}'
            self.filesUniqueFile(5)        
                
    def createAllPlan(self):
        wb = openpyxl.Workbook()
        for n, name in enumerate(self.sheets):
            if n == 0:
                sheet = wb.active 
                sheet.title = name
            else:
                sheet = wb.create_sheet(name)
            dataSheet = self.df[name].fillna('')            
            dataSheet = dataSheet.to_dict(orient='records')
            for data in dataSheet:
                values = list(data.values())
                sheet.append(values)
        wb.save(self.fileOut)
    
    def csvCsv(self, mode):
        allLines = []
        for data in self.dataFile:
            try:
                newData = [str(item).encode('utf-8-sig').decode('utf-8-sig') for item in data]
            except: 
                newData = [str(item).encode('ISO-8859-1').decode('utf-8-sig') for item in data]
            allLines.append(newData)
            newAllLines = []
            for line in allLines:
                newLine = []
                for lin in line:
                    newLin = lin.replace('ï»¿', '')
                    newLine.append(newLin)
                newAllLines.append(newLine)
            with open(self.fileOut, 'w', newline='', encoding='utf-8-sig') as recordCsv:
                writerCsv = csv.writer(recordCsv)
                writerCsv.writerows(newAllLines)
        if mode == 1:
            return self.fileOut
            
    def xlsXlsxOdsManyFormats(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1] 
            try:
                dfAllDict = allDfs[f]   
                self.dfAll = dfAllDict[0]
                cnt = 1
                for name, df in self.dfAll.items():
                    self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                    self.df = dfAllDict[cnt].fillna('')
                    self.renameHead()
                    if self.index == 2:
                        self.df.to_csv(self.fileOut, sep='\t', index=False,  encoding='utf-8-sig')
                    elif self.index == 5: 
                        self.df.to_json(self.fileOut, orient='records', date_format='iso', indent=4, force_ascii=True)
                    else:
                        self.df.to_csv(self.fileOut, index=False, encoding='utf-8-sig')
                    cnt += 1
                    self.bytesFiles()
            except:
                if self.index == 5:
                    self.dfAll = pd.read_excel(self.dataFile, sheet_name=None)
                else:
                    self.dfAll = pd.read_excel(self.dataFile, sheet_name=None, engine=self.engine)
                for name, df in self.dfAll.items():
                    self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}_{name}.{self.ext}'
                    self.df = df.fillna('')
                    self.renameHead()
                    if self.index == 2:
                        self.df.to_csv(self.fileOut, sep='\t', index=False,  encoding='utf-8-sig')
                    elif self.index == 5: 
                        try:
                            self.df.to_json(self.fileOut, orient='records', date_format='iso', indent=4, force_ascii=True)
                        except:
                            self.df.to_json(self.fileOut, orient='split', date_format='iso', indent=4, force_ascii=True)
                    else:
                        self.df.to_csv(self.fileOut, index=False, encoding='utf-8-sig')
                    self.bytesFiles() 
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            self.filesUniqueFile(0)
    
    def filesUniqueFile(self, category):
        try:
            if category in [0, 1, 2, 3, 6, 7]:
                for name, df in self.dfAll.items():
                    self.df = df.fillna('')
                    self.df.insert(loc=0, column=name, value=name, allow_duplicates=True)
                    self.df = df.fillna('')
                    self.df = self.df.astype(str)
                combinedDf = pd.concat(self.dfAll.values(), ignore_index=True)
                if category == 0:
                    if self.index == 2:
                        combinedDf.to_csv(self.fileOut, sep='\t', index=False,  encoding='utf-8-sig')
                    elif self.index == 5: 
                        try:
                            combinedDf.to_json(self.fileOut, orient='records', date_format='iso', indent=4, force_ascii=True)
                        except:
                            combinedDf.to_json(self.fileOut, orient='split', date_format='iso', indent=4, force_ascii=True)
                    else:
                        combinedDf.to_csv(self.fileOut, index=False, encoding='utf-8-sig')
                elif category == 1:
                    try:
                        combinedDf.to_excel(self.fileOut, index=False, na_rep='')
                    except: 
                        combinedDf.to_excel(self.fileOut, index=False, engine=self.engine, na_rep='')
                elif category == 2:
                    htmlStr = combinedDf.to_html(index=False, border=1, classes='dataframe')
                    with open(self.fileOut, 'w', encoding='utf-8') as f:
                        f.write(htmlStr)
                elif category == 3:
                    if self.index == 0:
                        htmlStr = combinedDf.to_html()
                        with open(self.fileOut, 'w', encoding='utf-8') as f:
                            f.write(htmlStr)
                    elif self.index == 4:
                        yamlData = {}
                        yamlData[name] = combinedDf.to_dict(orient='records')
                        with open(self.fileOut, 'w', encoding='utf-8') as outfile:
                            yaml.dump(yamlData, outfile, sort_keys=False, indent=4, allow_unicode=True)
                    elif self.index == 6:
                        htmlTable = combinedDf.to_html(index=True, border=1, classes='dataframe', na_rep='', 
                                               justify='center', show_dimensions=True)
                        external = configExternal(None)
                        xhtmlContent = external.includeXhtml(htmlTable)
                        self.fileOut = f'{self.nameFile}.{self.ext}'
                        with open(self.fileOut, 'w', encoding='utf-8') as f:
                            f.write(xhtmlContent)                
                    elif self.index == 7:
                        tomData = {}
                        tomData[name] = combinedDf.to_dict(orient='records')
                        with open(self.fileOut, 'w', encoding='utf-8') as outfile:
                            yaml.dump(tomData, outfile, sort_keys=False, indent=4, allow_unicode=True)
                elif category == 6:
                    self.df = combinedDf
                    self.prepairePdf()
                elif category == 7:
                    self.fileProv = f'{self.nameFile}_prov.ods'
                    with pd.ExcelWriter(self.fileProv, engine="odf") as writer:
                        df.to_excel(writer, index=False, sheet_name=name)
                    self.fileOds = self.fileProv
                    pyexcel.save_as(file_name=self.fileOds, dest_file_name=self.fileOut)
            elif category == 4:
                mergedDoc = Document()
                nDocs = len(self.docs)
                for index, file in enumerate(self.docs):
                    subDoc = Document(file)
                    if index < nDocs-1:
                       subDoc.add_page_break()
                    for element in subDoc.element.body:
                        mergedDoc.element.body.append(element)
                mergedDoc.save(self.fileOut)
            elif category == 5:
                self.outOds = self.fileOut
                self.sheets = list(self.df.keys())
                self.fileOut = f'{self.nameFile}_prov.xls'
                self.createAllPlan()
                self.df = pd.read_excel(self.fileOut, sheet_name=None)
                self.fileOut = f'{self.nameFile}_prov.xlsx'
                with pd.ExcelWriter(self.fileOut, engine='openpyxl') as writer:
                    for sheet_name, df in self.df.items():
                        df.to_excel(writer, sheet_name=sheet_name, index=False)            
                dataXlsx = get_data(self.fileOut) 
                save_data(self.outOds, dataXlsx) 
                self.fileOut = self.outOds
            self.bytesFiles()
        except:
            pass
    
    def csvXlsx(self):
        for f, file in enumerate(self.files):
            self.nameFile = file[0]
            self.dataFile = file[1]
            self.code = file[2]
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            self.sheetName = 'aba_única'
            self.csvPlan()
            self.bytesFiles()
            
    def csvPlan(self):
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = self.sheetName        
        for data in self.dataFile:
            try:
                newData = [str(item).encode('ISO-8859-1').decode('utf-8-sig') for item in data]
            except: 
                newData = [str(item).encode('utf-8-sig').decode('utf-8-sig') for item in data]
            sheet.append(newData)
        wb.save(self.fileOut)
    
    def csvTsv(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            self.bytesFiles()
               
    def csvDf(self, exprFile):
        file = self.files[0]
        self.nameFile = file[0]
        self.dataFile = file[1]
        self.fileOut = self.nameFile
        self.fileOut = self.csvCsv(1)
        self.df = pd.read_csv(self.fileOut, encoding='utf-8-sig').fillna('')
        self.df = self.df.astype(str)
        self.renameHead()
        self.exprFile = exprFile
        self.expr = ''
        self.exprLine, self.exprCol, self.exprCells = [w for w in range(3)]
        self.returnRowCol()
        self.place.markdown(f'{self.exprFile} (*{self.expr}*)') 
        st.write(self.df)
        
    def returnRowCol(self):
        objs = []
        lines, cols = self.df.shape
        objects = {0:[self.exprLine, lines, 'linha', 'linhas'], 
                   1:[self.exprCol, cols, 'coluna', 'colunas'], 
                   2:[self.exprCells, lines*cols, 'célula', 'células']}
        for obj, ects in objects.items():
            exprObj = ects[0]
            valObj = ects[1]
            if valObj <= 1:
               obj = f'{valObj} {ects[2]}'
            elif all([valObj >= 2, valObj <= 999]):
                obj = f'{valObj} {ects[3]}'
            else:
                valForm = f'{valObj:,.0f}'.replace(',', '.')
                obj = f'{valForm} {ects[3]}'
            objs.append(obj)
        self.expr = ', '.join(objs[:-1])
        self.expr += f' e {objs[-1]}'
        self.place = st.empty()
        self.place.write('')
        
    def xlsXslxOdsDf(self, pos, exprFile, keyEngine):
        self.exprFile = exprFile
        file = self.files[0]
        self.nameFile = file[0]
        self.dataFile = file[1]
        allDfs.clear()
        allDfs.setdefault(pos, [])
        self.dfAll = pd.read_excel(self.dataFile, sheet_name=None, engine=keyEngine)
        allDfs[pos].append(self.dfAll)
        nAbas = len(self.dfAll.items())
        if nAbas <= 1:
            exprAbas = f'🔰 {nAbas} aba {exprFile}'
        else:
            exprAbas = f'🔰 {nAbas} abas {exprFile}'
        cnt = 1
        place = st.empty()
        place.write('')
        place.markdown(exprAbas)
        for name, df in self.dfAll.items():
            self.df = pd.read_excel(self.dataFile, sheet_name=name, engine=keyEngine)
            self.df = self.df.fillna('')
            self.renameHead()
            allDfs[pos].append(self.df)
            self.expr = ''
            self.exprLine, self.exprCol, self.exprCells = [w for w in range(3)]
            self.returnRowCol()
            st.markdown(f'#️⃣  {cnt} de {nAbas} <-> :red[**{name}**] (*{self.expr}*)') 
            st.write(self.df)
            cnt += 1
    
    def csvDocx(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            doc = Document()
            with open(self.fileOut, 'r', encoding='utf-8-sig') as f:
                csvReader = csv.reader(f)
                headers = next(csvReader)
                num_cols = len(headers)
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Table Grid'
                table.page_width = Inches(11.7) 
                table.page_height = Inches(8.5)                 
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for row in csvReader:
                    row_cells = table.add_row().cells
                    for i, cell_data in enumerate(row):
                        row_cells[i].text = cell_data
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
                doc.save(self.fileOut)
                self.bytesFiles()
                
    def csvYaml(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            df = pd.read_csv(self.fileOut).fillna('')
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            dataDict = df.to_dict(orient='records')
            with open(self.fileOut, 'w', encoding='utf-8') as outfile:
                yaml.dump(dataDict, outfile, sort_keys=False, indent=4, allow_unicode=True)
            self.bytesFiles()
            
    def csvXhtml(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            df = pd.read_csv(self.fileOut)
            htmlTable = df.to_html(index=False, border=1, classes='dataframe')
            external = configExternal(None)
            xhtmlContent = external.includeXhtml(htmlTable)
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            with open(self.fileOut, 'w', encoding='utf-8') as f:
                f.write(xhtmlContent)
            self.bytesFiles()
                    
    def csvJson(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            df = pd.read_csv(self.fileOut).fillna('')
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            df.to_json(self.fileOut, orient='records', date_format='iso', indent=4, force_ascii=True)
            self.bytesFiles()
            
    def csvToml(self):
        for f, file in enumerate(self.files):
            dataAll = []
            self.file = file
            self.prepaireCsv()
            with open(self.fileOut, mode='r', encoding='utf-8') as csvFile:
                csvReader = csv.DictReader(csvFile)
                for row in csvReader:
                    processedRow = {}
                    for key, value in row.items():
                        if value is None or value == '':
                            processedRow[key] = None 
                        elif value.lower() in [True, False]:
                            processedRow[key] = value.lower() == 'true'
                        else:
                            try:
                                processedRow[key] = int(value)
                            except ValueError:
                                try:
                                    processedRow[key] = float(value)
                                except ValueError:
                                    processedRow[key] = value 
                    dataAll.append(processedRow)
            self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
            with open(self.fileOut, mode='w', encoding='utf-8') as tomlFile:
                toml.dump({'dados completos': dataAll}, tomlFile)
            self.bytesFiles()
    
    def csvTxt(self):
        try:
            for f, file in enumerate(self.files):
                self.file = file
                self.prepaireCsv()
                df = pd.read_csv(self.fileOut).fillna('')
                self.fileOut = f'{str(f+1).zfill(5)}_{self.nameFile}.{self.ext}'
                df.to_csv(self.fileOut, index=False, header=True)
                self.bytesFiles()
        except:
            pass
    
    def csvPdf(self):
        for f, file in enumerate(self.files):
            self.file = file
            self.prepaireCsv()
            self.df = pd.read_csv(self.fileOut).fillna('')
            self.df = self.df.astype(str)
            self.fileOut = f'{self.nameFile}.{self.ext}' 
            self.prepairePdf()             
            self.bytesFiles()

    def prepairePdf(self):
        pdf = FPDF(orientation='L', unit='in', format='A4')
        pdf.add_page()
        pdf.set_font("Arial", size=4)
        self.tam = pdf.w 
        self.listCols = []
        self.col_width = pdf.w / (len(self.df.columns) + 1)
        self.defineColPdf()
        row_height = 0.3
        for c, col in enumerate(self.df.columns):
            col_width = self.listCols[c]
            pdf.cell(col_width, row_height, col, border=1, align='C')
        pdf.ln()
        for index, row in self.df.iterrows():
            for d, data in enumerate(row.values):
                col_width = self.listCols[d]
                try:
                    dataStr=str(data)
                    pdf.cell(col_width, row_height, dataStr, border=1, align='C')
                except:
                    self.data = str(data)
                    self.rectifyData()
                    pdf.cell(col_width, row_height, self.data, border=1, align='C')
            pdf.ln()
        pdf.output(self.fileOut)
    
    def rectifyData(self):
        padSymb = r"[^\w\s]" 
        symbols = re.findall(padSymb, self.data)
        excts = ['/', ',', '.', ';', '_']
        for lt in self.data:
            if all([lt in symbols, lt not in excts]):
                self.data = self.data.replace(lt, '').strip()                 
    
    def defineColPdf(self):
        sizeCols = {}
        for c, col in enumerate(self.df.columns):
            sizeCols.setdefault(c, [])
            sizeCols[c].append(len(str(col))) 
        for index, row in self.df.iterrows():
            for d, data in enumerate(row.values):
                sizeCols[d].append(len(str(data)))            
        maxCols = []
        keyCols = list(sizeCols.keys())
        for key in keyCols: 
            maxCols.append(max(sizeCols[key]))
        for size in maxCols: 
            divide = size/sum(maxCols)*self.tam
            if divide > self.col_width:
                self.listCols.append(divide*0.75)
            else:
                self.listCols.append(self.col_width*0.90)
        
    def prepaireCsv(self):
        self.nameFile = self.file[0]
        self.dataFile = self.file[1]
        self.fileOut = f'{self.nameFile}_new.csv'
        self.csvCsv(0)
    
    def filesAllDf(self):
        listFile = self.files[0]
        listName = list(listFile.keys())
        listValues = list(listFile.values())
        listOriginal = list(map(lambda name: f'{sepFile}'.join(name.split(sepFile)[:-1]), listName))        
        isRepet = False
        if self.opt == -1:
            listSit = []
            for value in listValues:
                if value == 1:
                    listSit.append('não repetido')
                else:
                    difValue = value - 1
                    if difValue >= 1:
                        isRepet = True                    
                    if difValue == 1:
                        exprDif = 'vez'
                    else:
                        exprDif = 'vezes'
                    listSit.append(f'repetido {difValue} {exprDif}')
            if isRepet:
                st.markdown(f'📚 Arquivos redundantes')
                repls = ['Recusar❓']
                dfRepls = {'Arquivos repetidos':repls}
                optRepl = st.dataframe(dfRepls, 
                                       selection_mode="single-row",
                                       on_select="rerun",
                                       height='stretch', 
                                       hide_index=True)
                selRepl = optRepl.selection.rows
                if selRepl:
                    st.session_state[replDown[0]] = True
                else:
                    st.session_state[replDown[0]] = False
            dfDict = {'nome original': listOriginal, 
                      'pseudônimo': listName,                       
                      'número de arquivos': listValues, 
                      'situação': listSit}
            sumAll = sum(list(listFile.values()))
            st.markdown(f'🖥️ Detalhes e visualização ({sumAll})')
            df = pd.DataFrame(dfDict)
            df = df.astype(str)            
            event = st.dataframe(df,
                                 selection_mode="single-row",
                                 on_select="rerun",
                                 width=720, 
                                 height='stretch')
            selInd = event.selection.rows
            if selInd:
                ilocIndex = selInd[0]
                fileRow = df.iloc[[ilocIndex], 1].tolist()
                fileSelDf.clear()
                fileSelDf.append(fileRow[0])                
        else:
            if self.opt == -2:
                nameNotRep = [listName[v] for v in range(len(listValues)) if listValues[v] >= 1] 
                nameOrig = [f'{sepFile}'.join(listName[v].split(sepFile)[:-1]) for v in range(len(listValues)) if listValues[v] >= 1]      
                dfDict = {'nome original': nameOrig,
                          'pseudônimo': nameNotRep}                           
            else:
                nameRep = [listName[v] for v in range(len(listValues)) if listValues[v] > 1]
                nameOrig = [f'{sepFile}'.join(listName[v].split(sepFile)[:-1]) for v in range(len(listValues)) if listValues[v] > 1]
                numRep = [listValues[v]-1 for v in range(len(listValues)) if listValues[v] > 1]
                dfDict = {'nome original': nameOrig,
                          'pseudônimo': nameRep, 
                          'número de repetições': numRep}
            df = pd.DataFrame(dfDict)
            df = df.astype(str)
            st.dataframe(df)  

    def bytesFiles(self):
        output = BytesIO()
        with open(self.fileOut, 'rb') as arquivo:
            docRead = arquivo.read()
        zips = (self.fileOut, docRead)
        self.filesZip.append(zips) 
        self.nFiles += 1
    
    def renameHead(self):
        head = {}
        for col in self.df.columns:
            if col.lower().find('unnamed') >= 0:
                head[col] = ''
            else:
                head[col] = col
        self.df.rename(columns=head, inplace=True)
        
    def downZip(self):
        with zipfile.ZipFile(self.nameZip, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file in self.filesZip:
                nameFile = file[0]
                dataFile = file[1]
                zipf.writestr(nameFile, dataFile)
                
class configExternal():
    def __init__(self, *args):
        self.args = args
        pass
        
    def configCss(self):
        with open('configCss.css') as f:
            css = f.read()
        st.markdown(f'<style>{css}</style>', unsafe_allow_html=True) 

    def configSelect(self):
        st.markdown(f"""
        <style>
            .st-e4 {{max-width: {self.args[0]}px !important;}} 
        </style>
        """, unsafe_allow_html=True)        
        
    def configJson(self):
        js = f"""
            <script>
                var element = window.parent.document.getElementById("{self.args[0]}");
                if (element) {{
                    element.scrollIntoView({{behavior: 'instant', block: 'start'}});
                }}
            </script>
            """
        components.html(js, height=0, width=0)  

    def includeXhtml(self, htmlTable):
        xhtmlContent = f"""<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Strict//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd">
            <html xmlns="http://www.w3.org/1999/xhtml" lang="en" xml:lang="en">
            <head>
                <title>CSV to XHTML Table</title>
                <meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
            </head>
            <body>
                {htmlTable}
            </body>
            </html>"""
        return xhtmlContent
        
class main():
    def __init__(self):          
        st.set_page_config(initial_sidebar_state="collapsed", layout="wide")
        self.typeExt = ['CSV', 'XLS', 'XLSX', 'ODS']
        self.engine = ['', 'openpyxl', 'xlrd', 'openpyxl', 'odf']
        nIni = len(self.typeExt)
        self.typeExt.insert(0, '')
        colType, colUpload = st.columns([12, 17], width='stretch')
        self.keyUp = 'zero'
        self.keyFile = 'typeFile'
        self.keyRep = 'keyRep'
        with colType:
            with st.container(border=4, key='contType', gap='small', height="stretch"):
                colStart, colIco = st.columns([0.5, 20], vertical_alignment='top')
                st.markdown('<div id="start"></div>', unsafe_allow_html=True)
                colIco.markdown('❇️ Seleção de tipo + arrastamento/escolha de arquivos', unsafe_allow_html=True, 
                            text_alignment='center')
                self.typeFile = st.selectbox(f'📂 Tipos de arquivo ({nIni})', self.typeExt,
                                                      help=f'Selecionar a extensão desejada. Para reiniciar, ' 
                                                            'escolher a linha em branco. Por padrão, arquivos repetidos'
                                                            'são admitidos. Para evitar isso, acione o botão "Informações" e marque '
                                                            'a opção "Recusar".') 
                if not self.typeFile: 
                    upDisabled = True
                    repDisabled = True
                    self.typeStr = ''
                else:
                    self.loc = self.typeExt.index(self.typeFile)
                    upDisabled = False
                    repDisabled = False
                    self.typeStr = f':red[**{self.typeFile}**]'
                    st.space(size="small")  
                self.upLoad = st.file_uploader(f'☰ Arraste/escolha um ou múltiplos arquivos {self.typeStr}.', 
                                               type=self.typeFile, accept_multiple_files=True, key=self.keyUp, 
                                               disabled=upDisabled, 
                                               help='É integrado de todos os arquivos selecionados, mesmo que se repitam. ' 
                                                    'No momento do download, o usuário poderá acionar o botão "Informações" e '
                                                    'marcar a opção "Recusar". Só é ativado quando houver extensão escolhida.') 
        with colUpload:  
            try:
                self.files = list(set([f'{file.name}{sepFile}{file.size}' for file in self.upLoad]))
                for file in self.upLoad: 
                    allNames.append(f'{file.name}{sepFile}{file.size}')        
            except:
                self.files = [] 
                allNames.clear()
            if not self.typeFile:
                self.configImageEmpty(4)
            if replDown[0] not in st.session_state:
                st.session_state[replDown[0]] = False
            self.repFile = st.session_state[replDown[0]]
            if self.typeFile:
                self.ext = self.typeFile.lower()
                with st.container(border=4, key='contUpload', gap='small', height='content', 
                                  vertical_alignment='center'):
                    self.nUpLoads = len(self.upLoad)
                    match self.loc:
                        case 1:
                            self.exts = {self.engine[1]: ['xls', 'xlsx', 'html'], 'odf': ['ods'], 'tsv': ['tsv'], 
                                         'doc': ['docx'], 'yaml': ['yaml'], 'json': ['json'], 'xhtml': ['xhtml'],
                                         'toml': ['toml'], 'txt': ['txt'], 'pdf': ['pdf']}  
                        case 2: 
                            self.exts = {self.engine[2]: ['csv'.lower(), 'xlsx', 'html'], 'odf': ['ods'], 'tsv': ['tsv'], 
                                         'doc': ['docx'], 'yaml': ['yaml'], 'json': ['json'], 'xhtml': ['xhtml'],
                                         'toml': ['toml'], 'txt': ['txt'], 'pdf': ['pdf']}  
                        case 3: 
                            self.exts = {self.engine[3]: ['csv'.lower(), 'xls', 'html'], 'odf': ['ods'], 'tsv': ['tsv'], 
                                         'doc': ['docx'], 'yaml': ['yaml'], 'json': ['json'], 'xhtml': ['xhtml'],
                                         'toml': ['toml'], 'txt': ['txt'], 'pdf': ['pdf']} 
                        case 4: 
                            self.exts = {self.engine[3]: ['xls', 'xlsx', 'html'], 'odf': ['csv'], 'tsv': ['tsv'], 
                                         'doc': ['docx'], 'yaml': ['yaml'], 'json': ['json'], 'xhtml': ['xhtml'],
                                         'toml': ['toml'], 'txt': ['txt'], 'pdf': ['pdf']} 
                    self.newTypes = []
                    self.segregateTypes()
                    self.disableds = ['disabled' + str(w) for w in range(len(self.newTypes))]
                    if self.nUpLoads == 0:
                        self.setSessionState(True)
                    else:
                        self.setSessionState(False)
                    typeLow = self.typeFile.lower()
                    self.strFunc = ['Converter um ou mais arquivos', 'Convertendo']
                    self.stripe = f':red[**{self.typeFile.upper()}**]'
                    with st.container(border=None, key='contOne', gap='small', height='stretch', 
                                      vertical_alignment='center'):
                        nFiles = len(self.files)
                        if nFiles <= 0:
                            titleSel = f'Arquivo selecionado ({nFiles})'
                        else:
                            titleSel = f'Arquivos selecionados ({nFiles})'
                        if nFiles > 0:
                            opts = sorted(self.files)
                            opts.insert(0, '')
                        else:
                            opts = []
                        buttOne, buttTwo, buttThree, buttFour, buttFive, buttSix = ['' for i in range(6)]
                        buttSeven, buttEight, buttNine, buttTen, buttEleven, buttTwelve = ['' for i in range(6)]
                        self.allButtons = [buttOne, buttTwo, buttThree, buttFour, buttFive, buttSix, 
                                           buttSeven, buttEight, buttNine, buttTen, buttEleven, buttTwelve]
                        colOne, colTwo, colThree = st.columns(spec=3, width='stretch')
                        colFour, colFive, colSix = st.columns(spec=3, width='stretch')
                        colSeven, colEight, colNine = st.columns(spec=3, width='stretch')
                        colTen, colEleven, colTwelve = st.columns(spec=3, width='stretch')
                        self.colsButts = {0: [0, colOne, ':material/sync_alt:'], 1: [1, colTwo, ':material/swap_horiz:'], 
                                          2: [2, colThree, ':material/table_convert:'], 3: [3, colFour, ':material/transform:'], 
                                          4: [4, colFive, ':material/convert_to_text:'], 5: [5, colSix, ':material/edit_arrow_up:'], 
                                          6: [6, colSeven, ':material/business_messages:'], 7: [7, colEight, ':material/edit_document:'], 
                                          8: [8, colNine, ':material/edit_square:'], 9: [9, colTen, ':material/edit_note:'], 
                                          10: [10, colEleven, ':material/box_edit:'], 11: [11, colTwelve, ':material/contract_edit:']}
                        for b, buttObj in enumerate(self.allButtons):
                            buttObj = self.setButtons(self.colsButts[b])
                        indOpt = {0: [0, 0], 1: [0, 1], 2: [0, 2], 3: [1, 0], 4: [2, 0], 5: [3, 0],
                                  6: [4, 0], 7: [5, 0], 8: [6, 0], 9: [7, 0], 10: [8, 0], 11: [9, 0]}
                    if self.upLoad:
                        filesAll, filesRep, nNotRep, nRep, exprLoad, exprNotRep, exprRep = self.allNotRep()
                        nUps = self.nUpLoads
                        textUp, textNotRep, textRep = ('', '', '')
                        dctSize = {0: [nUps, textUp, '', ''], 
                                   1: [nNotRep, textNotRep, ' não repetido', ' não repetidos'], 
                                   2: [nRep, textRep, ' repetido', ' repetidos']}
                        for dct, size in dctSize.items():
                            if size[0] == 1:
                               size[1] = f'do arquivo{size[2]}'
                            else:
                               size[1] = f'dos arquivos{size[3]}'     
                        self.files.insert(0, '')
                        indExt = allExts.index(self.ext)
                        try:
                            self.filesRead = [] 
                            self.segregateFiles()                                                     
                        except:
                            pass
                        with st.container(border=None, key='contRepNo', gap='small', height='content', 
                                          vertical_alignment='bottom'):
                            colTotal, colNotRep, colRep = st.columns(spec=3, width='stretch', 
                                                                     vertical_alignment='center')
                            with colTotal.popover(f'Informações ({nUps})', icon='ℹ️', width='stretch', 
                                                  help=f'Abre tela com detalhes e possibilidade de visualização {dctSize[0][1]}.'): 
                                downOrDfFiles([filesAll, self.files[1:], filesRep], None, None, self.ext, -1, None, None)
                                if len(fileSelDf) > 0: 
                                    self.elem = fileSelDf[0]
                                    nameElem = f'{sepFile}'.join(self.elem.split(sepFile)[:-1])
                                    try:
                                        self.organizeDf()
                                    except Exception as error: 
                                        place = st.empty()
                                        place.write('')
                                        objMens = messages(None, None, None)
                                        objMens.mensOperation(f'🚫 Houve o seguinte erro\n *:blue-background[{error}]*.')                                        
                                st.text('')   
                            with colNotRep.popover(f'{exprNotRep} ({nNotRep})', icon='👍', width='stretch', 
                                                   help=f'Abre tela com detalhes {dctSize[1][1]}.'):
                                st.markdown(f'✒️ Sem redundância ({nNotRep})', width=720)
                                downOrDfFiles([filesAll, self.files[1:], filesRep], None, None, None, -2, None, None)
                            if nRep == 0:
                                disabledRep = True
                            else:
                                disabledRep = False
                            with colRep.popover(f'{exprRep} ({nRep})', icon='👎', width='stretch', disabled=disabledRep, 
                                                help=f'Abre tela com detalhes {dctSize[2][1]}.'):
                                st.markdown(f'✒️ Com redundância ({nRep})', width=720)
                                downOrDfFiles([filesAll, self.files[1:], filesRep], None, None, None, -3, None, None)  
                        if any(self.allButtons):
                            if self.loc in [1, 2, 3, 4]: 
                                ind = self.allButtons.index(True) 
                                self.expr = f'{self.strFunc[1]} {self.nUpLoads} do formato {self.stripe} para o formato {self.newTypes[ind]}...'
                                self.index, self.opt = indOpt[ind]
                                self.keys = list(self.exts.keys())
                                self.key = self.keys[self.index]
                                self.values = self.exts[self.key]
                                self.ext = self.values[self.opt]
                                try:
                                    place = st.empty()
                                    place.write('')
                                    self.preInvoke()  
                                except Exception as error:  
                                    objMens = messages(None, None, None)
                                    objMens.mensOperation(f'⚠️ Houve o seguinte erro\n *:yellow-background[{error}]*.')
                            
    def preInvoke(self):
        if st.session_state[replDown[0]]:
            self.cutFilesRep()
        with st.spinner(self.expr):
            downOrDfFiles(self.filesRead, self.index, self.key, self.ext, self.opt, 
                          self.typeFile, self.typeExt[1:])
    
    def cutFilesRep(self):
        mylist = [(file[0], file[-1]) for file in self.filesRead]
        allRep = []
        for my in mylist:
            locs = [i for i, item in enumerate(mylist) if item == my]
            if locs not in allRep:
                allRep.append(locs)
        allRep = [rep[0] for rep in allRep]
        self.filesRead = [self.filesRead[w] for w in allRep]
    
    def allNotRep(self):
        filesAll = {}
        filesRep = []
        for file in self.upLoad:
            prefix = f'{file.name}{sepFile}{file.size}'
            if prefix in self.files: 
                filesAll.setdefault(prefix, 0)
                filesAll[prefix] += 1
                nNames = filesAll[prefix]                
                if nNames > 1:
                    filesRep.append(prefix)
        nNotRep = len(self.files)
        nRep = len(filesRep)                              
        exprLoad = self.singPlural(self.nUpLoads, 'arquivo não repetido', 'arquivos não repetidos')
        exprNotRep = self.singPlural(nNotRep, 'não repetido', 'não repetidos')
        exprRep = self.singPlural(nRep, 'repetido', 'repetidos')  
        return(filesAll, filesRep, nNotRep, nRep, exprLoad, exprNotRep, exprRep)  
    
    def organizeDf(self):
        nameFile = self.elem
        self.ext = self.typeFile.lower()
        self.pos = allNames.index(nameFile)
        self.filesReadDf = [] 
        self.segregateDf()
        with st.spinner('Aguarde a exibição do arquivo na tela...'):
            objDown = downOrDfFiles(self.filesReadDf, None, None, self.ext, None, None, None)
            if self.loc == 1:
                exprFile = f'📋 Aba do arquivo :red[**{nameFile}**]'
                objDown.csvDf(exprFile)   
            elif self.loc in [2, 3, 4]:
                exprFile = f' de :red[**{nameFile}**]'
                engine = self.engine[self.loc]
                objDown.xlsXslxOdsDf(self.pos, exprFile, engine)
                    
    def singPlural(self, *args):
        if args[0] <= 1: 
            expr = args[1]
        else:
            expr = args[2]
        return expr
    
    def setButtons(self, elems):
        n = elems[0]
        col = elems[1]
        ico = elems[2]
        butt = f'butt{n}'
        labelButt = f'{self.stripe} para {self.newTypes[n]}'
        self.allButtons[n] = col.button(label=labelButt, key=butt, width='stretch', 
                                        icon=ico, 
                                        disabled=st.session_state[self.disableds[n]],
                                        help=f'{self.strFunc[0]} {self.stripe} para {self.newTypes[n]}.')
    
    def segregateTypes(self):
        listTypes = list(self.exts.values())
        for tipo in listTypes:
            self.newTypes += tipo        
        self.newTypes = [f':red[**{new.upper()}**]' for new in self.newTypes]
    
    def configImageEmpty(self, border):
        with st.container(border=border, key='contZero', gap='small'):
            st.markdown(f'0️⃣  seleção de tipo e/ou arquivo', text_alignment='center') 
            st.image('zero.jpg') 
    
    def setSessionState(self, state):
        for disabled in self.disableds:
            if disabled not in st.session_state:
                st.session_state[disabled] = True 
            else:
                st.session_state[disabled] = state
              
    def segregateFiles(self):
        filesFind = {}
        if self.loc == 1:
            for upLoad in self.upLoad: 
                nameGlobal = upLoad.name
                nameFile, ext = os.path.splitext(nameGlobal)
                nameSize = f'{nameFile}_{upLoad.size}'
                filesFind.setdefault(nameGlobal, 0)
                if nameGlobal in self.files:
                    filesFind[nameGlobal] += 1
                if filesFind[nameGlobal] > 1:
                    continue
                dataBytes = upLoad.getvalue()
                dataString = dataBytes.decode('ISO-8859-1')
                self.fileMemory = io.StringIO(dataString)
                sep = self.detectSep()
                readerCsv = csv.reader(self.fileMemory, delimiter=sep)
                joinNameRead = (nameFile, readerCsv, nameSize)
                self.filesRead.append(joinNameRead)
        elif self.loc in [2, 3, 4]: 
            for upLoad in self.upLoad: 
                nameGlobal = upLoad.name
                nameFile, ext = os.path.splitext(nameGlobal)
                nameSize = f'{nameFile}_{upLoad.size}'
                filesFind.setdefault(nameGlobal, 0)
                if nameGlobal in self.files:
                    filesFind[nameGlobal] += 1
                if filesFind[nameGlobal] > 1:
                    continue
                bytesExcel = BytesIO(upLoad.read())
                joinNameRead = (nameFile, bytesExcel, nameSize)
                self.filesRead.append(joinNameRead)
            
    def segregateDf(self):        
        if self.loc == 1:
            for u, upLoad in enumerate(self.upLoad):
                if u == self.pos:
                    nameGlobal = upLoad.name
                    nameFile, ext = os.path.splitext(nameGlobal)
                    nameSize = f'{nameFile}_{upLoad.size}'
                    dataBytes = upLoad.getvalue()
                    dataString = dataBytes.decode('ISO-8859-1')
                    self.fileMemory = io.StringIO(dataString)
                    sep = self.detectSep()
                    readerCsv = csv.reader(self.fileMemory, delimiter=sep)
                    joinNameRead = (nameFile, readerCsv, nameSize, sep)
                    self.filesReadDf.append(joinNameRead)
                    break         
        elif self.loc in [2, 3, 4]: 
            for u, upLoad in enumerate(self.upLoad):
                if u == self.pos:
                    nameGlobal = upLoad.name
                    nameFile, ext = os.path.splitext(nameGlobal)
                    nameSize = f'{nameFile}_{upLoad.size}'
                    joinNameRead = (nameFile, upLoad, nameSize)
                    self.filesReadDf.append(joinNameRead)
                    break                    
    
    def detectSep(self):
        lines = 1024*10
        sample = self.fileMemory.read(lines)
        self.fileMemory.seek(0)
        dialect = csv.Sniffer().sniff(sample)
        return dialect.delimiter
            
if __name__ == '__main__':
    global sepFile, fileSelDf, allNames
    global allDfs, allEngines, allExts
    global replDown
    sepFile = '_'
    fileSelDf = []
    allNames = []
    allDfs = {}
    allEngines = ['openpyxl', 'xlrd', 'odf']
    allExts = ['csv', 'xls', 'xlsx', 'ods']
    replDown = ['selRepl']
    external = configExternal(None)
    external.configCss()
    main()

